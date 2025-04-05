import streamlit as st
import requests
import base64
import json
from supabase import create_client
from datetime import datetime
import re
import tempfile
import os
import time
import google.generativeai as genai
import docx
from docx.shared import Pt
import io

# Set page config FIRST (must be the first Streamlit command)
st.set_page_config(
    page_title="Learning Observer", 
    layout="wide",
    page_icon="📝"
)

# Helper function to parse secrets from GitHub secret
def parse_secrets(secret_content):
    """Parse secrets from the single secret string"""
    secrets = {}
    for line in secret_content.split('\n'):
        if '=' in line:
            key, value = line.split('=', 1)
            secrets[key.strip()] = value.strip()
    return secrets

# Load secrets from GitHub secret or Streamlit secrets
try:
    # Try to get the single secret from GitHub
    secret_content = os.getenv('SECRET_TRY', '')
    if not secret_content:
        # Fall back to Streamlit secrets
        secret_content = st.secrets.get("SECRET_TRY", "")
    
    secrets = parse_secrets(secret_content)
    
    # Extract individual keys
    SUPABASE_URL = secrets.get("SUPABASE_URL", "")
    SUPABASE_KEY = secrets.get("SUPABASE_KEY", "")
    GOOGLE_API_KEY = secrets.get("GOOGLE_API_KEY", "")
    ASSEMBLYAI_API_KEY = secrets.get("ASSEMBLYAI_API_KEY", "")
    OCR_API_KEY = secrets.get("OCR_API_KEY", "")
    GROQ_API_KEY = secrets.get("GROQ_API_KEY", "")
    
except Exception as e:
    st.error(f"Error loading secrets: {str(e)}")
    SUPABASE_URL = ""
    SUPABASE_KEY = ""
    GOOGLE_API_KEY = ""
    ASSEMBLYAI_API_KEY = ""
    OCR_API_KEY = ""
    GROQ_API_KEY = ""

# Initialize Supabase client
@st.cache_resource
def init_supabase():
    return create_client(SUPABASE_URL, SUPABASE_KEY)

supabase = init_supabase()

# Configure Google AI API
genai.configure(api_key=GOOGLE_API_KEY)

class ObservationExtractor:
    def __init__(self):
        self.ocr_api_key = OCR_API_KEY
        self.groq_api_key = GROQ_API_KEY
        self.gemini_api_key = GOOGLE_API_KEY

    def image_to_base64(self, image_file):
        """Convert image file to base64 string"""
        return base64.b64encode(image_file.read()).decode('utf-8')

    def extract_text_with_ocr(self, image_file):
        """Extract text from image using OCR.space API"""
        try:
            # Get file extension
            file_type = image_file.name.split('.')[-1].lower()
            if file_type == 'jpeg':
                file_type = 'jpg'

            # Convert image to base64
            base64_image = self.image_to_base64(image_file)
            base64_image_with_prefix = f"data:image/{file_type};base64,{base64_image}"

            # Prepare request payload
            payload = {
                'apikey': self.ocr_api_key,
                'language': 'eng',
                'isOverlayRequired': False,
                'iscreatesearchablepdf': False,
                'issearchablepdfhidetextlayer': False,
                'OCREngine': 2,  # Better for handwriting
                'detectOrientation': True,
                'scale': True,
                'base64Image': base64_image_with_prefix
            }

            # Send request to OCR API
            response = requests.post(
                'https://api.ocr.space/parse/image',
                data=payload,
                headers={'apikey': self.ocr_api_key}
            )

            response.raise_for_status()
            data = response.json()

            # Process response
            if not data.get('ParsedResults') or len(data['ParsedResults']) == 0:
                error_msg = data.get('ErrorMessage', 'No parsed results returned')
                raise Exception(f"OCR Error: {error_msg}")

            parsed_result = data['ParsedResults'][0]
            if parsed_result.get('ErrorMessage'):
                raise Exception(f"OCR Error: {parsed_result['ErrorMessage']}")

            extracted_text = parsed_result['ParsedText']

            if not extracted_text or not extracted_text.strip():
                raise Exception("No text was detected in the image")

            return extracted_text

        except Exception as e:
            st.error(f"OCR Error: {str(e)}")
            raise

    def process_with_groq(self, extracted_text):
        """Process extracted text with Groq AI"""
        try:
            # Original detailed prompt
            system_prompt = """You are an AI assistant for a learning observation system. Extract and structure information from the provided observation sheet text.

The observation sheets typically have the following structure:
- Title (usually "The Observer")
- Student information (Name, Roll Number/ID)
- Date and Time information
- Core Observation Section with time slots
- Teaching content for each time slot
- Learning details (what was learned, tools used, etc.)

Format your response as JSON with the following structure:
{
  "studentName": "Student's name if available, otherwise use the title of the observation",
  "studentId": "Student ID or Roll Number",
  "className": "Class name or subject being taught",
  "date": "Date of observation",
  "observations": "Detailed description of what was learned",
  "strengths": ["List of strengths observed in the student"],
  "areasOfDevelopment": ["List of areas where the student needs improvement"],
  "recommendations": ["List of recommended actions for improvement"]
}

For observations, provide full detailed descriptions like:
"The student learned how to make maggi from their mom through in-person mode, including all steps from boiling water to adding spices"

Be creative in extracting information based on context."""

            # Send request to Groq API
            response = requests.post(
                'https://api.groq.com/openai/v1/chat/completions',
                headers={
                    'Authorization': f'Bearer {self.groq_api_key}',
                    'Content-Type': 'application/json'
                },
                json={
                    "model": "llama-3.3-70b-versatile",
                    "messages": [
                        {
                            "role": "system",
                            "content": system_prompt
                        },
                        {
                            "role": "user",
                            "content": f"Extract and structure the following text from an observation sheet: {extracted_text}"
                        }
                    ],
                    "temperature": 0.2,
                    "response_format": {"type": "json_object"}
                }
            )

            response.raise_for_status()
            data = response.json()

            # Extract the JSON content
            ai_response = data['choices'][0]['message']['content']
            return json.loads(ai_response)

        except Exception as e:
            st.error(f"Groq API Error: {str(e)}")
            raise

    def transcribe_with_assemblyai(self, audio_file):
        """Transcribe audio using AssemblyAI API"""
        if not ASSEMBLYAI_API_KEY:
            return "Error: AssemblyAI API key is not configured. Please add it to your secrets."
            
        # Set up the API headers
        headers = {
            "authorization": ASSEMBLYAI_API_KEY,
            "content-type": "application/json"
        }
        
        # Upload the audio file
        try:
            st.write("Uploading audio file...")
            upload_response = requests.post(
                "https://api.assemblyai.com/v2/upload",
                headers={"authorization": ASSEMBLYAI_API_KEY},
                data=audio_file.getvalue()
            )
            
            if upload_response.status_code != 200:
                return f"Error uploading audio: {upload_response.text}"
            
            upload_url = upload_response.json()["upload_url"]
            
            # Request transcription
            st.write("Processing transcription...")
            transcript_request = {
                "audio_url": upload_url,
                "language_code": "en"
            }
            
            transcript_response = requests.post(
                "https://api.assemblyai.com/v2/transcript",
                json=transcript_request,
                headers=headers
            )
            
            if transcript_response.status_code != 200:
                return f"Error requesting transcription: {transcript_response.text}"
            
            transcript_id = transcript_response.json()["id"]
            
            # Poll for completion
            status = "processing"
            progress_bar = st.progress(0)
            while status != "completed" and status != "error":
                polling_response = requests.get(
                    f"https://api.assemblyai.com/v2/transcript/{transcript_id}",
                    headers=headers
                )
                
                if polling_response.status_code != 200:
                    return f"Error checking transcription status: {polling_response.text}"
                
                polling_data = polling_response.json()
                status = polling_data["status"]
                
                if status == "completed":
                    progress_bar.progress(100)
                    return polling_data["text"]
                elif status == "error":
                    return f"Transcription error: {polling_data.get('error', 'Unknown error')}"
                
                # Update progress
                progress = polling_data.get("percent_done", 0)
                if progress:
                    progress_bar.progress(progress / 100.0)
                time.sleep(2)
            
            return "Error: Transcription timed out or failed."
        except Exception as e:
            return f"Error during transcription: {str(e)}"

    def generate_report_from_text(self, text_content, user_info):
        """Generate a structured report from text using Google Gemini"""
        prompt = f"""
        Based on this text from a student observation, create a detailed observer report following the Observer Report format.
        
        TEXT CONTENT:
        {text_content}
        
        FORMAT REQUIREMENTS:
        
        1. Daily Activities Overview - Extract and categorize the student's daily activities into:
           - Morning activities
           - Afternoon activities
           - Evening activities
           - Night activities (if mentioned)
        
        2. Learning Moments & Reflections - Identify:
           - New skills or knowledge the student gained
           - Interesting observations or experiences
           - Any self-reflection shared
        
        3. Thinking Process - Assess:
           - Approach to new information (curious/skeptical/analytical/accepting)
           - Logical thinking (strong/moderate/needs improvement)
           - Problem-solving skills (effective/developing/needs guidance)
           - Creativity and imagination (high/moderate/low)
           - Decision-making style (confident/hesitant/experimental)
           - Any unique perspectives or ideas
        
        4. Communication Skills & Thought Clarity - Evaluate:
           - Confidence level (low/medium/high)
           - Clarity of thought (clear/slightly clear/confused)
           - Participation & engagement (active/moderate/passive)
           - Sequence of explanation (well-structured/partially structured/unstructured)
        
        5. General Behavior & Awareness - Note:
           - Behavior (polite/calm/energetic/distracted)
           - General awareness (aware/partially aware/unaware)
        
        6. Observer's Comments - Add any relevant observations
        
        7. Summary for Parents - Write a brief paragraph summarizing the session
        
        Use the exact section titles and format as above. For items that cannot be determined from the text, use "Not enough information" rather than making assumptions.
        """
        
        try:
            # Configure the model - using Gemini Pro for most comprehensive responses
            model = genai.GenerativeModel('gemini-1.5-pro-002')
            
            # Generate content with Gemini
            response = model.generate_content([
                {"role": "user", "parts": [{"text": prompt}]}
            ])
            
            # Extract the content from the response
            report_content = response.text
            
            # Add user information to the report
            complete_report = f"""Date: {user_info['session_date']}
    Student Name: {user_info['student_name']}
    Observer Name: {user_info['observer_name']}
    Session Duration: {user_info['session_start']} - {user_info['session_end']}

    {report_content}

    Name of Observer: {user_info['observer_name']}
    """
            return complete_report
        except Exception as e:
            return f"Error generating report: {str(e)}"

    def create_word_document(self, report_content):
        """Create a Word document from the report content with proper formatting"""
        doc = docx.Document()
        
        # Add title 
        title = doc.add_heading('The Observer Report', 0)
        
        # Clean up markdown formatting
        report_content = report_content.replace('**', '')
        
        # Add report content, parsing the sections
        lines = report_content.split('\n')
        section_pattern = re.compile(r'^\d+\.\s+(.+)')
        subheading_pattern = re.compile(r'^\*\s*(.+):\*\s*(.+)')
        list_item_pattern = re.compile(r'^\*\s+(.+)')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Header information (Date, Name, etc.)
            if line.startswith(('Date:', 'Student Name:', 'Observer Name:', 'Session Duration:', 'Name of Observer:')):
                p = doc.add_paragraph()
                p.add_run(line).bold = True
            
            # Section heading (e.g., "1. Daily Activities Overview")
            elif section_match := section_pattern.match(line):
                doc.add_heading(line, level=1)
            
            # Subheading with bold prefix (e.g., "* Morning activities: Woke up early")
            elif subheading_match := subheading_pattern.match(line):
                p = doc.add_paragraph()
                prefix = subheading_match.group(1)
                content = subheading_match.group(2)
                p.add_run(f"{prefix}: ").bold = True
                p.add_run(content)
            
            # List item
            elif list_match := list_item_pattern.match(line):
                content = list_match.group(1)
                p = doc.add_paragraph(content, style='List Bullet')
            
            # Regular paragraph
            else:
                doc.add_paragraph(line)
        
        # Save to a BytesIO object
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes

# Main App
def main():
    # Initialize extractor
    extractor = ObservationExtractor()

    # Session state
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False
    if 'username' not in st.session_state:
        st.session_state.username = ""
    if 'user_info' not in st.session_state:
        st.session_state.user_info = {
            'student_name': '',
            'observer_name': '',
            'session_date': datetime.now().strftime('%d/%m/%Y'),
            'session_start': '',
            'session_end': ''
        }
    if 'audio_transcription' not in st.session_state:
        st.session_state.audio_transcription = ""
    if 'report_generated' not in st.session_state:
        st.session_state.report_generated = None
    if 'show_edit_transcript' not in st.session_state:
        st.session_state.show_edit_transcript = False
    if 'processing_mode' not in st.session_state:
        st.session_state.processing_mode = None  # 'ocr' or 'audio'

    # Login Page (accepts any credentials)
    if not st.session_state.authenticated:
        st.title("Learning Observer Login")

        with st.form("login_form"):
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            submitted = st.form_submit_button("Login")

            if submitted:
                st.session_state.authenticated = True
                st.session_state.username = username
                st.rerun()

        return

    # Main Application
    st.title(f"Welcome, {st.session_state.username}")

    if st.button("Logout"):
        st.session_state.authenticated = False
        st.session_state.username = ""
        st.rerun()

    # Sidebar for user information
    with st.sidebar:
        st.subheader("Session Information")
        
        # Get user input
        st.session_state.user_info['student_name'] = st.text_input("Student Name:", value=st.session_state.user_info['student_name'])
        st.session_state.user_info['observer_name'] = st.text_input("Observer Name:", value=st.session_state.user_info['observer_name'])
        st.session_state.user_info['session_date'] = st.date_input("Session Date:").strftime('%d/%m/%Y')
        
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.user_info['session_start'] = st.time_input("Start Time:").strftime('%H:%M')
        with col2:
            st.session_state.user_info['session_end'] = st.time_input("End Time:").strftime('%H:%M')

    # Choose processing mode
    st.subheader("Select Processing Mode")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("OCR Mode (Image Upload)"):
            st.session_state.processing_mode = "ocr"
            st.session_state.audio_transcription = ""
            st.session_state.report_generated = None
    with col2:
        if st.button("Audio Mode (Recording Upload)"):
            st.session_state.processing_mode = "audio"
            st.session_state.audio_transcription = ""
            st.session_state.report_generated = None

    if st.session_state.processing_mode == "ocr":
        st.info("OCR Mode: Upload an image of an observation sheet")
        
        # File uploader for images
        uploaded_file = st.file_uploader("Upload Observation Sheet", type=["jpg", "jpeg", "png"])
        
        if uploaded_file is not None:
            st.image(uploaded_file, caption="Uploaded Image", use_column_width=True)
            
            if st.button("Process Observation"):
                with st.spinner("Processing..."):
                    try:
                        # Step 1: OCR
                        extracted_text = extractor.extract_text_with_ocr(uploaded_file)
                        
                        # Step 2: Groq AI processing
                        structured_data = extractor.process_with_groq(extracted_text)
                        
                        # Generate report from structured data
                        observations_text = structured_data.get("observations", "")
                        if observations_text:
                            report = extractor.generate_report_from_text(observations_text, st.session_state.user_info)
                            st.session_state.report_generated = report
                            
                            # Save to Supabase
                            supabase.table('observations').insert({
                                "username": st.session_state.username,
                                "student_name": structured_data.get("studentName", ""),
                                "student_id": structured_data.get("studentId", ""),
                                "class_name": structured_data.get("className", ""),
                                "date": structured_data.get("date", ""),
                                "observations": observations_text,
                                "strengths": json.dumps(structured_data.get("strengths", [])),
                                "areas_of_development": json.dumps(structured_data.get("areasOfDevelopment", [])),
                                "recommendations": json.dumps(structured_data.get("recommendations", [])),
                                "timestamp": datetime.now().isoformat(),
                                "filename": uploaded_file.name,
                                "full_data": json.dumps(structured_data)  # Store complete data
                            }).execute()
                            
                            st.success("Data processed and saved successfully!")
                        else:
                            st.error("No observations found in the extracted data")
                            
                    except Exception as e:
                        st.error(f"Processing error: {str(e)}")

    elif st.session_state.processing_mode == "audio":
        st.info("Audio Mode: Upload an audio recording of an observation session")
        
        # File uploader for audio
        uploaded_file = st.file_uploader("Choose an audio file", 
                                       type=["wav", "mp3", "m4a", "mpeg", "mpg", "ogg", "flac", "aac", "wma", "aiff"])
        
        if uploaded_file is not None:
            # Display audio player
            st.audio(uploaded_file)
            
            # Check if user information is filled
            if (st.session_state.user_info['student_name'] and 
                st.session_state.user_info['observer_name'] and 
                st.session_state.user_info['session_start'] and 
                st.session_state.user_info['session_end']):
                
                # One-click process button
                if st.button("Process & Generate Report"):
                    if not ASSEMBLYAI_API_KEY:
                        st.error("AssemblyAI API key is missing. Please configure it in your secrets.")
                    else:
                        # First, transcribe the audio
                        with st.spinner("Step 1/2: Transcribing audio..."):
                            transcript = extractor.transcribe_with_assemblyai(uploaded_file)
                            st.session_state.audio_transcription = transcript
                        
                        # Then, generate the report
                        with st.spinner("Step 2/2: Generating report with Gemini AI..."):
                            report = extractor.generate_report_from_text(transcript, st.session_state.user_info)
                            st.session_state.report_generated = report
                            
                            # Save to Supabase
                            supabase.table('observations').insert({
                                "username": st.session_state.username,
                                "student_name": st.session_state.user_info['student_name'],
                                "student_id": "",
                                "class_name": "",
                                "date": st.session_state.user_info['session_date'],
                                "observations": transcript,
                                "strengths": json.dumps([]),
                                "areas_of_development": json.dumps([]),
                                "recommendations": json.dumps([]),
                                "timestamp": datetime.now().isoformat(),
                                "filename": uploaded_file.name,
                                "full_data": json.dumps({
                                    "transcript": transcript,
                                    "report": report
                                })
                            }).execute()
            else:
                st.warning("Please fill in all the session information in the sidebar.")

    # Option to edit transcript (only shown if transcript exists)
    if st.session_state.audio_transcription:
        if st.button("Edit Transcription" if not st.session_state.show_edit_transcript else "Hide Editor"):
            st.session_state.show_edit_transcript = not st.session_state.show_edit_transcript
        
        if st.session_state.show_edit_transcript:
            st.subheader("Edit Transcription")
            edited_transcript = st.text_area("You can edit the transcript below:", 
                                            value=st.session_state.audio_transcription, 
                                            height=200)
            
            # Update the transcription if edited
            if edited_transcript != st.session_state.audio_transcription:
                st.session_state.audio_transcription = edited_transcript
            
            # Regenerate report with edited transcript
            if st.button("Regenerate Report with Edited Transcript"):
                with st.spinner("Generating report with Gemini AI..."):
                    report = extractor.generate_report_from_text(
                        st.session_state.audio_transcription,
                        st.session_state.user_info
                    )
                    st.session_state.report_generated = report

    # Show report if generated
    if st.session_state.report_generated:
        st.subheader("Generated Report")
        
        # Display the report in the app
        st.markdown(st.session_state.report_generated)
        
        # Create Word document for download
        docx_file = extractor.create_word_document(st.session_state.report_generated)
        
        # Download button
        student_name = st.session_state.user_info['student_name'].replace(" ", "_")
        date = st.session_state.user_info['session_date'].replace("/", "-")
        filename = f"Observer_Report_{student_name}_{date}.docx"
        
        st.download_button(
            label="Download as Word Document",
            data=docx_file,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Sample data for testing
    with st.expander("Use sample data for testing"):
        if st.button("Load sample transcript"):
            sample_transcript = """
            So today I'd like to talk about my day as a student. This morning I woke up around 7 AM and spent some time reviewing my math notes before school. I've been trying to understand calculus better. I had breakfast with my family and then caught the bus at 8.

            At school, I had four classes. The science class was really interesting because we did an experiment about chemical reactions. I noticed that when we mixed certain compounds, the color changed immediately, which made me wonder about the electron transfer happening. During lunch, I sat with my friends and we talked about the upcoming science fair.

            In the afternoon, I had my debate club meeting. We're preparing for a competition next month. I'm working on improving my rebuttal techniques. I think I'm getting better at thinking on my feet when challenged with counterarguments.

            After school, I went to the library for about an hour to work on my history project. I'm researching the industrial revolution and its impact on social structures. It's fascinating how technology changes society.

            When I got home around 5 PM, I spent some time playing piano. I'm learning a new piece by Mozart. It's challenging but I've been breaking it down into smaller sections to practice. After dinner with my family, I spent about two hours on homework, mostly math problems and some reading for English class.

            Before bed, I watched a short documentary about space exploration on YouTube. I've always been curious about astronomy. I think that's everything important from my day.
            """
            st.session_state.audio_transcription = sample_transcript
            st.experimental_rerun()

    # Footer
    st.markdown("---")
    st.markdown("The Observer Report Generator")


if __name__ == "__main__":
    main()
